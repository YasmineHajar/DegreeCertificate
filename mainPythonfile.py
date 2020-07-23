import subprocess, os, platform, webbrowser
import docx, xlrd
import numpy as np
from docx2python import docx2python

def replace_string2(filename,i,dirName):
    doc = docx.Document(filename)
    firstname=i[0]
    lastname = i[1]
    for p in doc.paragraphs:
        if 'First' in p.text:
            text = p.text.replace('First', firstname)
            style = p.style
            p.text = text
            p.style = style
        if 'Last' in p.text:
            #print ('SEARCH FOUND!!')
            text = p.text.replace('Last', lastname)
            style = p.style
            p.text = text
            p.style = style
    newfilename=lastname+"_"+firstname+".docx"
    filepath = dirName + "/" + newfilename
    doc.save(filepath)
    return 1


def readlistnamefromword(filename):
    list = []
    doc = docx.Document(filename)
    tables=doc.tables
    for row in tables[0].rows:
        newname=(str(row.cells[0].text)+" " +str(row.cells[1].text)+" "+str(row.cells[2].text))
        print (newname)
        list.append(newname)
    return list

def readlistnamefromexcel(filename,n):
    list1 = []
    workbook = xlrd.open_workbook(filename)
    worksheet = workbook.sheet_by_name('Sheet1')
    #worksheet = workbook.sheet_by_index(0)
    m=n
    data1 = worksheet.col_values(m)
    list1 = [item for item in data1 if item is not None]
    return list1



list1=readlistnamefromexcel("NameList.xlsx",0)
list2=readlistnamefromexcel("NameList.xlsx",1)
#print (list1)
#print (list2)

list=np.vstack((list1, list2)).T
#list=np.column_stack(list1, list2)
l=len(list1)
print(list)
j=0

# Create target Directory if don't exist
dirName = 'ListName'
if not os.path.exists(dirName):
    os.mkdir(dirName)
    print("Directory " , dirName ,  " Created ")
else:
    #os.rmdir(dirName)
    #shutil.rmtree(dirName)
    print("Directory " , dirName ,  " Replaced")

for i in list:
    replace_string2('DegreeForm.docx',i,dirName)

